"""
Word Document Generator for Stock Briefing Reports

브리핑 리포트를 Word 문서(.docx)로 생성하는 모듈입니다.

Usage:
    from docx_generator import create_briefing_report

    briefing_data = {
        'date': '2025-12-24',
        'title': '오늘의 화제 종목 TOP 5',
        'summary': '시장 요약...',
        'stocks': [...]
    }

    output_path = create_briefing_report(briefing_data, 'output/briefing.docx')
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
import logging

logger = logging.getLogger(__name__)


class BriefingReportGenerator:
    """브리핑 리포트 생성기"""

    def __init__(self, title_font: str = 'Arial', body_font: str = 'Calibri',
                 title_size: int = 24, body_size: int = 11):
        """
        Args:
            title_font: 제목 폰트
            body_font: 본문 폰트
            title_size: 제목 크기
            body_size: 본문 크기
        """
        self.doc = Document()
        self.title_font = title_font
        self.body_font = body_font
        self.title_size = title_size
        self.body_size = body_size
        self._setup_styles()

    def _setup_styles(self):
        """문서 스타일 설정"""
        # Heading 1 스타일
        style = self.doc.styles['Heading 1']
        font = style.font
        font.name = self.title_font
        font.size = Pt(self.title_size)
        font.bold = True
        font.color.rgb = RGBColor(0, 51, 102)  # 다크 블루

        # Heading 2 스타일
        style = self.doc.styles['Heading 2']
        font = style.font
        font.name = self.title_font
        font.size = Pt(18)
        font.bold = True
        font.color.rgb = RGBColor(37, 99, 235)  # 블루

        # Heading 3 스타일
        style = self.doc.styles['Heading 3']
        font = style.font
        font.name = self.title_font
        font.size = Pt(14)
        font.bold = True

        # Normal 스타일
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.body_font
        font.size = Pt(self.body_size)

    def add_title_section(self, title: str, date: str):
        """제목 섹션 추가"""
        # 제목
        title_para = self.doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 날짜
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"📅 {date}")
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(100, 100, 100)

        # 구분선
        self.doc.add_paragraph('=' * 60)

    def add_summary_section(self, summary: str):
        """요약 섹션 추가"""
        self.doc.add_heading('📊 시장 요약', level=2)
        para = self.doc.add_paragraph(summary)
        para.style = 'Body Text'
        self.doc.add_paragraph()

    def add_stock_section(self, stock: Dict, rank: int, include_chart: bool = True):
        """
        종목 섹션 추가

        Args:
            stock: 종목 데이터
            rank: 순위
            include_chart: 차트 이미지 포함 여부
        """
        # 종목 헤더
        self.doc.add_heading(
            f"{rank}. {stock.get('symbol', 'N/A')} - {stock.get('name', 'Unknown')}",
            level=3
        )

        # 종목 정보 표
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        # 현재가
        cells = table.rows[0].cells
        cells[0].text = '현재가'
        cells[1].text = f"${stock.get('price', 0):.2f}"

        # 등락률
        cells = table.rows[1].cells
        cells[0].text = '등락률'
        change = stock.get('change_percent', 0)
        cells[1].text = f"{change:+.2f}%"

        # 등락률 색상 설정
        for paragraph in cells[1].paragraphs:
            for run in paragraph.runs:
                if change >= 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # 녹색
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
                run.font.bold = True

        # 거래량
        cells = table.rows[2].cells
        cells[0].text = '거래량'
        cells[1].text = f"{stock.get('volume', 0):,}"

        # 뉴스 요약
        cells = table.rows[3].cells
        cells[0].text = '뉴스 요약'
        cells[1].text = stock.get('news_summary', 'N/A')

        # 차트 이미지 삽입 (있는 경우)
        if include_chart:
            chart_path = stock.get('chart_image_path')
            if chart_path and Path(chart_path).exists():
                try:
                    self.doc.add_paragraph()
                    self.doc.add_picture(chart_path, width=Inches(5))
                    caption = self.doc.add_paragraph()
                    caption.add_run(f"그림: {stock.get('symbol')} 주가 차트").italic = True
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to insert chart image: {e}")

        self.doc.add_paragraph()

    def add_footer_section(self):
        """푸터 섹션 추가"""
        self.doc.add_paragraph('=' * 60)

        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = (
            f"생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"출처: Yahoo Finance, Exa News API\n"
            f"🤖 Generated with Claude Code - While You Were Sleeping"
        )
        footer_run = footer.add_run(footer_text)
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)

    def save(self, output_path: str) -> str:
        """
        문서 저장

        Args:
            output_path: 저장 경로

        Returns:
            저장된 파일 경로
        """
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        logger.info(f"Report saved to: {output_path}")
        return output_path


def create_briefing_report(
    briefing_data: Dict,
    output_path: str,
    include_charts: bool = True
) -> str:
    """
    브리핑 리포트 생성

    Args:
        briefing_data: 브리핑 데이터 딕셔너리
            - date: 날짜 (YYYY-MM-DD)
            - title: 제목
            - summary: 요약
            - stocks: 종목 리스트
        output_path: 저장 경로
        include_charts: 차트 이미지 포함 여부

    Returns:
        생성된 문서 경로

    Example:
        briefing_data = {
            'date': '2025-12-24',
            'title': '오늘의 화제 종목 TOP 5',
            'summary': '시장 요약...',
            'stocks': [
                {
                    'symbol': 'AAPL',
                    'name': 'Apple Inc.',
                    'price': 150.25,
                    'change_percent': 2.5,
                    'volume': 50000000,
                    'news_summary': 'Apple announces...',
                    'chart_image_path': 'path/to/chart.png'  # 선택
                }
            ]
        }
        create_briefing_report(briefing_data, 'output/briefing.docx')
    """
    try:
        generator = BriefingReportGenerator()

        # 제목 섹션
        generator.add_title_section(
            title=briefing_data.get('title', '일일 주식 브리핑'),
            date=briefing_data.get('date', datetime.now().strftime('%Y-%m-%d'))
        )

        # 요약 섹션
        generator.add_summary_section(
            summary=briefing_data.get('summary', '')
        )

        # 종목 섹션
        generator.doc.add_heading('🔥 화제 종목 TOP 5', level=2)
        for i, stock in enumerate(briefing_data.get('stocks', [])[:5], 1):
            generator.add_stock_section(stock, rank=i, include_chart=include_charts)

        # 푸터 섹션
        generator.add_footer_section()

        # 저장
        return generator.save(output_path)

    except Exception as e:
        logger.error(f"Failed to create briefing report: {e}")
        raise


def create_stock_comparison_report(stocks: List[Dict], output_path: str) -> str:
    """
    종목 비교 리포트 생성

    Args:
        stocks: 종목 리스트
        output_path: 저장 경로

    Returns:
        생성된 문서 경로

    Example:
        stocks = [
            {'symbol': 'AAPL', 'name': 'Apple Inc.', 'price': 150.0, 'change_percent': 2.5, 'volume': 50000000},
            {'symbol': 'MSFT', 'name': 'Microsoft', 'price': 380.0, 'change_percent': 1.8, 'volume': 28000000}
        ]
        create_stock_comparison_report(stocks, 'output/comparison.docx')
    """
    try:
        doc = Document()

        # 제목
        title = doc.add_heading('종목 비교 분석', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 비교 표 생성
        table = doc.add_table(rows=len(stocks) + 1, cols=6)
        table.style = 'Medium Grid 1 Accent 1'

        # 헤더
        headers = ['순위', '티커', '종목명', '현재가', '등락률', '거래량']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

        # 데이터 입력
        for i, stock in enumerate(stocks, 1):
            row = table.rows[i]
            row.cells[0].text = str(i)
            row.cells[1].text = stock.get('symbol', 'N/A')
            row.cells[2].text = stock.get('name', 'N/A')
            row.cells[3].text = f"${stock.get('price', 0):.2f}"

            # 등락률 (색상 적용)
            change = stock.get('change_percent', 0)
            change_cell = row.cells[4]
            change_cell.text = f"{change:+.2f}%"
            for paragraph in change_cell.paragraphs:
                for run in paragraph.runs:
                    if change >= 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    else:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    run.font.bold = True

            row.cells[5].text = f"{stock.get('volume', 0):,}"

        # 저장
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        logger.info(f"Comparison report saved to: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Failed to create comparison report: {e}")
        raise


def create_news_summary_report(news_items: List[Dict], output_path: str) -> str:
    """
    뉴스 요약 리포트 생성

    Args:
        news_items: 뉴스 아이템 리스트
        output_path: 저장 경로

    Returns:
        생성된 문서 경로

    Example:
        news_items = [
            {
                'title': 'News Title',
                'date': '2025-12-24',
                'source': 'Bloomberg',
                'summary': 'News summary...',
                'related_stocks': ['AAPL', 'MSFT'],
                'url': 'https://...'
            }
        ]
        create_news_summary_report(news_items, 'output/news.docx')
    """
    try:
        doc = Document()

        # 제목
        doc.add_heading('📰 주요 뉴스 요약', level=1)

        for i, news in enumerate(news_items, 1):
            # 뉴스 번호 및 제목
            doc.add_heading(f"{i}. {news.get('title', 'No Title')}", level=2)

            # 메타 정보 (날짜, 출처)
            meta = doc.add_paragraph()
            meta_run = meta.add_run(
                f"📅 {news.get('date', 'N/A')} | 🔗 {news.get('source', 'N/A')}"
            )
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(100, 100, 100)

            # 뉴스 요약
            summary = doc.add_paragraph(news.get('summary', ''))
            summary.style = 'Body Text'

            # 관련 종목
            if 'related_stocks' in news and news['related_stocks']:
                related = doc.add_paragraph()
                related_run = related.add_run(
                    f"관련 종목: {', '.join(news.get('related_stocks', []))}"
                )
                related_run.font.italic = True
                related_run.font.size = Pt(10)

            # URL
            if 'url' in news:
                url_para = doc.add_paragraph()
                url_para.add_run('링크: ')
                url_run = url_para.add_run(news.get('url', ''))
                url_run.font.color.rgb = RGBColor(0, 0, 255)
                url_run.font.underline = True

            doc.add_paragraph()  # 뉴스 간 간격

        # 저장
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        logger.info(f"News summary report saved to: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Failed to create news summary report: {e}")
        raise


if __name__ == "__main__":
    # 간단한 테스트
    logging.basicConfig(level=logging.INFO)

    test_data = {
        'date': '2025-12-24',
        'title': '테스트 브리핑',
        'summary': '이것은 테스트 브리핑입니다.',
        'stocks': [
            {
                'symbol': 'AAPL',
                'name': 'Apple Inc.',
                'price': 150.25,
                'change_percent': 2.5,
                'volume': 50000000,
                'news_summary': 'Apple announces new product...'
            }
        ]
    }

    output = create_briefing_report(test_data, 'output/reports/test_briefing.docx')
    print(f"Test report created: {output}")
