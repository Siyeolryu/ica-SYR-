"""
DOCX Report Generator Example

브리핑 리포트를 Word 문서로 생성하는 예제 코드입니다.

사용법:
    python example_generator.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from typing import Dict, List


class BriefingReportGenerator:
    """브리핑 리포트 생성기"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """문서 스타일 설정"""
        # Heading 1 스타일
        style = self.doc.styles['Heading 1']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(24)
        font.bold = True
        font.color.rgb = RGBColor(0, 51, 102)  # 다크 블루

        # Heading 2 스타일
        style = self.doc.styles['Heading 2']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(18)
        font.bold = True
        font.color.rgb = RGBColor(37, 99, 235)  # 블루

        # Heading 3 스타일
        style = self.doc.styles['Heading 3']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(14)
        font.bold = True

        # Normal 스타일
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def add_title_section(self, title: str, date: str):
        """제목 섹션 추가"""
        # 제목
        title_para = self.doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 날짜
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"📅 {date}")
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(100, 100, 100)

        # 구분선
        self.doc.add_paragraph('=' * 60)

    def add_summary_section(self, summary: str):
        """요약 섹션 추가"""
        self.doc.add_heading('📊 시장 요약', level=2)
        para = self.doc.add_paragraph(summary)
        para.style = 'Body Text'
        self.doc.add_paragraph()

    def add_stock_section(self, stock: Dict, rank: int):
        """종목 섹션 추가"""
        # 종목 헤더
        self.doc.add_heading(
            f"{rank}. {stock.get('symbol', 'N/A')} - {stock.get('name', 'Unknown')}",
            level=3
        )

        # 종목 정보 표
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        # 현재가
        cells = table.rows[0].cells
        cells[0].text = '현재가'
        cells[1].text = f"${stock.get('price', 0):.2f}"

        # 등락률
        cells = table.rows[1].cells
        cells[0].text = '등락률'
        change = stock.get('change_percent', 0)
        cells[1].text = f"{change:+.2f}%"

        # 등락률 색상 설정
        for paragraph in cells[1].paragraphs:
            for run in paragraph.runs:
                if change >= 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # 녹색
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
                run.font.bold = True

        # 거래량
        cells = table.rows[2].cells
        cells[0].text = '거래량'
        cells[1].text = f"{stock.get('volume', 0):,}"

        # 뉴스 요약
        cells = table.rows[3].cells
        cells[0].text = '뉴스 요약'
        cells[1].text = stock.get('news_summary', 'N/A')

        # 차트 이미지 삽입 (있는 경우)
        chart_path = stock.get('chart_image_path')
        if chart_path and Path(chart_path).exists():
            self.doc.add_paragraph()
            self.doc.add_picture(chart_path, width=Inches(5))
            caption = self.doc.add_paragraph()
            caption.add_run(f"그림: {stock.get('symbol')} 주가 차트").italic = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

    def add_footer_section(self):
        """푸터 섹션 추가"""
        self.doc.add_paragraph('=' * 60)

        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = (
            f"생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"출처: Yahoo Finance, Exa News API\n"
            f"🤖 Generated with Claude Code - While You Were Sleeping"
        )
        footer_run = footer.add_run(footer_text)
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)

    def save(self, output_path: str) -> str:
        """문서 저장"""
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        return output_path


def create_briefing_report(briefing_data: Dict, output_path: str) -> str:
    """
    브리핑 리포트 생성

    Args:
        briefing_data: 브리핑 데이터 딕셔너리
            - date: 날짜 (YYYY-MM-DD)
            - title: 제목
            - summary: 요약
            - stocks: 종목 리스트
        output_path: 저장 경로

    Returns:
        생성된 문서 경로
    """
    generator = BriefingReportGenerator()

    # 제목 섹션
    generator.add_title_section(
        title=briefing_data.get('title', '일일 주식 브리핑'),
        date=briefing_data.get('date', datetime.now().strftime('%Y-%m-%d'))
    )

    # 요약 섹션
    generator.add_summary_section(
        summary=briefing_data.get('summary', '')
    )

    # 종목 섹션
    generator.doc.add_heading('🔥 화제 종목 TOP 5', level=2)
    for i, stock in enumerate(briefing_data.get('stocks', [])[:5], 1):
        generator.add_stock_section(stock, rank=i)

    # 푸터 섹션
    generator.add_footer_section()

    # 저장
    return generator.save(output_path)


def create_stock_comparison_report(stocks: List[Dict], output_path: str) -> str:
    """
    종목 비교 리포트 생성

    Args:
        stocks: 종목 리스트
        output_path: 저장 경로

    Returns:
        생성된 문서 경로
    """
    doc = Document()

    # 제목
    title = doc.add_heading('종목 비교 분석', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 비교 표 생성
    table = doc.add_table(rows=len(stocks) + 1, cols=6)
    table.style = 'Medium Grid 1 Accent 1'

    # 헤더
    headers = ['순위', '티커', '종목명', '현재가', '등락률', '거래량']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # 데이터 입력
    for i, stock in enumerate(stocks, 1):
        row = table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = stock.get('symbol', 'N/A')
        row.cells[2].text = stock.get('name', 'N/A')
        row.cells[3].text = f"${stock.get('price', 0):.2f}"

        # 등락률 (색상 적용)
        change = stock.get('change_percent', 0)
        change_cell = row.cells[4]
        change_cell.text = f"{change:+.2f}%"
        for paragraph in change_cell.paragraphs:
            for run in paragraph.runs:
                if change >= 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.bold = True

        row.cells[5].text = f"{stock.get('volume', 0):,}"

    # 저장
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    """예제 실행"""
    print("DOCX Report Generator Example\n")

    # 예제 브리핑 데이터
    briefing_data = {
        'date': '2025-12-24',
        'title': '오늘의 화제 종목 TOP 5',
        'summary': (
            '오늘 미국 증시는 기술주 중심으로 상승세를 보였습니다. '
            '특히 AI 관련 주식들이 강세를 보이며 시장을 이끌었습니다. '
            '투자자들은 연말 랠리에 대한 기대감으로 적극적인 매수세를 보였습니다.'
        ),
        'stocks': [
            {
                'symbol': 'NVDA',
                'name': 'NVIDIA Corporation',
                'price': 495.50,
                'change_percent': 5.2,
                'volume': 45000000,
                'news_summary': 'NVIDIA의 새로운 AI 칩이 시장에서 큰 호응을 얻고 있습니다.'
            },
            {
                'symbol': 'TSLA',
                'name': 'Tesla Inc.',
                'price': 248.30,
                'change_percent': 3.8,
                'volume': 120000000,
                'news_summary': 'Tesla의 전기차 판매가 예상을 상회하며 주가가 급등했습니다.'
            },
            {
                'symbol': 'AAPL',
                'name': 'Apple Inc.',
                'price': 195.75,
                'change_percent': 2.1,
                'volume': 55000000,
                'news_summary': 'Apple의 신제품 출시 소식에 투자자들이 긍정적으로 반응했습니다.'
            },
            {
                'symbol': 'MSFT',
                'name': 'Microsoft Corporation',
                'price': 378.90,
                'change_percent': 1.9,
                'volume': 28000000,
                'news_summary': 'Microsoft의 클라우드 서비스 성장이 계속되고 있습니다.'
            },
            {
                'symbol': 'AMZN',
                'name': 'Amazon.com Inc.',
                'price': 155.20,
                'change_percent': -0.5,
                'volume': 42000000,
                'news_summary': 'Amazon은 일시적인 조정을 받았지만 여전히 강세를 유지하고 있습니다.'
            }
        ]
    }

    # 출력 디렉토리 설정
    output_dir = Path(__file__).parent.parent.parent.parent / 'backend' / 'output' / 'reports'
    output_dir.mkdir(parents=True, exist_ok=True)

    # 1. 표준 브리핑 리포트 생성
    print("1. Creating standard briefing report...")
    briefing_output = output_dir / 'example_briefing.docx'
    result1 = create_briefing_report(briefing_data, str(briefing_output))
    print(f"   ✓ Created: {result1}")

    # 2. 종목 비교 리포트 생성
    print("\n2. Creating stock comparison report...")
    comparison_output = output_dir / 'example_comparison.docx'
    result2 = create_stock_comparison_report(briefing_data['stocks'], str(comparison_output))
    print(f"   ✓ Created: {result2}")

    print("\n✅ All reports generated successfully!")
    print(f"\nOutput directory: {output_dir}")
    print("\nGenerated files:")
    print(f"  - {briefing_output.name}")
    print(f"  - {comparison_output.name}")


if __name__ == "__main__":
    main()
